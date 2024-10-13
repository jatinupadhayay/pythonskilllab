import pandas as pd
import tkinter as tk
from tkinter import ttk, filedialog, messagebox, simpledialog
from docx import Document
import random
from datetime import datetime, timedelta
subjects_df = pd.DataFrame()
rooms_df = pd.DataFrame()
faculty_df = pd.DataFrame()
generated_timetable = []

def read_excel_file(file_path):
    try:
        df = pd.read_excel(file_path)
        df.columns = df.columns.str.strip().str.lower()
        return df
    except Exception as e:
        messagebox.showerror("Error", f"Error reading file: {e}")
        return None


def load_subject_data():
    file_path = filedialog.askopenfilename(title="Select Subject Excel File", filetypes=[("Excel files", "*.xlsx *.xls")])
    if file_path:
        global subjects_df
        subjects_df = read_excel_file(file_path)
        if subjects_df is not None:
            display_subjects()


def display_subjects():
   
    for widget in subject_frame.winfo_children():
        widget.destroy()

    global subject_selected
    subject_selected = {}


    for year in subjects_df.columns:
        year_frame = ttk.LabelFrame(subject_frame, text=year.title())
        year_frame.pack(fill=tk.X, padx=5, pady=5)
        
        for subject in subjects_df[year].dropna().tolist():
            subject_var = tk.BooleanVar()
            checkbox = tk.Checkbutton(year_frame, text=subject, variable=subject_var)
            checkbox.pack(anchor='w')
            subject_selected[f"{year}_{subject}"] = subject_var


def add_subject():
    global subjects_df
    year = simpledialog.askstring("Input", "Enter Year:")
    subject_name = simpledialog.askstring("Input", "Enter Subject Name:")
    
    if year and subject_name:
        if year not in subjects_df.columns:
            subjects_df[year] = pd.Series(dtype=str)
        
        empty_index = subjects_df[year].isna().idxmax() if subjects_df[year].isna().any() else len(subjects_df)
        subjects_df.at[empty_index, year] = subject_name
        
        display_subjects()
        messagebox.showinfo("Success", f"Subject '{subject_name}' added to year '{year}'.")

def delete_subject():
    deleted_subjects = []
    for key, var in subject_selected.items():
        if var.get(): 
            parts = key.split('_')
            year = parts[0]  
            subject = '_'.join(parts[1:])  
            
           
            subjects_df[year].replace(subject, pd.NA, inplace=True)
            deleted_subjects.append(subject)
    
    if deleted_subjects:
        display_subjects()  
        messagebox.showinfo("Success", f"Deleted subjects: {', '.join(deleted_subjects)}")
    else:
        messagebox.showwarning("Warning", "No subject selected for deletion.")

def load_room_data():
    file_path = filedialog.askopenfilename(title="Select Room Excel File", filetypes=[("Excel files", "*.xlsx *.xls")])
    if file_path:
        global rooms_df
        rooms_df = read_excel_file(file_path)
        if rooms_df is not None:
            display_rooms()

def display_rooms():
    for widget in room_frame.winfo_children():
        widget.destroy()

    global room_selected
    room_selected = {}

    for room_type in rooms_df.columns:
        room_frame_type = ttk.LabelFrame(room_frame, text=room_type.title())
        room_frame_type.pack(fill=tk.X, padx=5, pady=5)
        
        for room in rooms_df[room_type].dropna().tolist():
            room_var = tk.BooleanVar()
            checkbox = tk.Checkbutton(room_frame_type, text=room, variable=room_var)
            checkbox.pack(anchor='w')
            room_selected[f"{room_type}_{room}"] = room_var

def add_room():
    room_type = simpledialog.askstring("Input", "Enter Room Type:")
    room_name = simpledialog.askstring("Input", "Enter Room Name:")
    
    if room_type and room_name:
        if room_type not in rooms_df.columns:
            rooms_df[room_type] = pd.Series(dtype=str)
        
        empty_index = rooms_df[room_type].isna().idxmax() if rooms_df[room_type].isna().any() else len(rooms_df)
        rooms_df.at[empty_index, room_type] = room_name
        
        display_rooms()
        messagebox.showinfo("Success", f"Room '{room_name}' added to type '{room_type}'.")

def delete_room():
    deleted_rooms = []
    for key, var in room_selected.items():
        if var.get(): 
            parts = key.split('_')
            room_type = parts[0]  
            room = '_'.join(parts[1:])
            
         
            rooms_df[room_type].replace(room, pd.NA, inplace=True)
            deleted_rooms.append(room)
    
    if deleted_rooms:
        display_rooms()  
        messagebox.showinfo("Success", f"Deleted rooms: {', '.join(deleted_rooms)}")
    else:
        messagebox.showwarning("Warning", "No room selected for deletion.")

def load_faculty_data():
    file_path = filedialog.askopenfilename(title="Select Faculty Excel File", filetypes=[("Excel files", "*.xlsx *.xls")])
    if file_path:
        global faculty_df
        faculty_df = read_excel_file(file_path)
        if faculty_df is not None:
            display_faculty()

def display_faculty():
    for widget in faculty_frame.winfo_children():
        widget.destroy()  

    global faculty_selected
    faculty_selected = {}

    for idx, row in faculty_df.iterrows():
        faculty_name = row.get('faculty name', 'N/A')
        occupation = row.get('occupation', 'N/A')
        experience = row.get('experience', 'N/A')  

        faculty_selected[faculty_name] = tk.BooleanVar()
        checkbox = tk.Checkbutton(faculty_frame, text=f"{faculty_name} ({occupation}, {experience} years)", variable=faculty_selected[faculty_name])
        checkbox.pack(anchor='w')

import tkinter as tk
from tkinter import ttk, messagebox, simpledialog
from datetime import datetime, timedelta
import random

import tkinter as tk
from tkinter import ttk, messagebox, simpledialog
from datetime import datetime, timedelta
import random

def generate_available_dates(start_date, num_days):
    return [(start_date + timedelta(days=i)).strftime("%Y-%m-%d") for i in range(num_days)]

def get_custom_time_slots(num_slots):
    time_slots = []
    for i in range(num_slots):
        time_slot = simpledialog.askstring("Time Slot", f"Enter time slot {i + 1} (e.g., 12:00 PM - 1:00 PM):")
        time_slots.append(time_slot)
    return time_slots

def get_exam_duration_and_slots():
    try:
        num_days = int(simpledialog.askstring("Exam Duration", "Enter exam duration (1, 2, or 3 days):"))

        if num_days not in [1, 2, 3]:
            raise ValueError("Invalid number of days.")

        exam_dates = []
        for i in range(num_days):
            exam_date = simpledialog.askstring("Exam Date", f"Enter date for day {i + 1} (YYYY-MM-DD):")
            exam_dates.append(exam_date)

        num_slots = int(simpledialog.askstring("Number of Slots", "Enter number of slots per day (1, 2, or 3):"))

        if num_slots not in [1, 2, 3]:
            raise ValueError("Invalid number of slots.")

       
        time_slots = get_custom_time_slots(num_slots)

        return exam_dates, time_slots
    except ValueError as ve:
        messagebox.showerror("Input Error", str(ve))
        return None, None

def check_conflicts(exams, new_exam):
    for exam in exams:
        # Check for faculty conflict
        if exam['date'] == new_exam['date'] and exam['time'] == new_exam['time']:
            if any(faculty in exam['faculty'] for faculty in new_exam['faculty']):
                return f"Faculty conflict: {', '.join(new_exam['faculty'])} already assigned to another exam at this time."

            # Check for room conflict
            if exam['room_number'] == new_exam['room_number']:
                return f"Room conflict: Room {new_exam['room_number']} already assigned to another exam at this time."

    # No conflicts detected, exam can be scheduled
    return "No conflicts"

def generate_timetable():
    if subjects_df.empty or rooms_df.empty or faculty_df.empty:
        messagebox.showwarning("Warning", "Please upload subject, room, and faculty files first.")
        return

    for widget in content_frame.winfo_children():
        if isinstance(widget, ttk.Treeview):
            widget.destroy()

    global generated_timetable
    generated_timetable = []  # List to hold the generated timetable
    exam_dates, time_slots = get_exam_duration_and_slots()
    if not exam_dates or not time_slots:
        return

    columns = ("Date", "Time", "Year", "Subject", "Room Number", "Building", "Faculty")
    timetable_tree = ttk.Treeview(content_frame, columns=columns, show='headings')

    for col in columns:
        timetable_tree.heading(col, text=col)
        timetable_tree.column(col, width=100, anchor='center')

    timetable_tree.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

    selected_faculty = [name for name, selected in faculty_selected.items() if selected.get()]

    if not selected_faculty:
        messagebox.showwarning("Warning", "No faculty members selected!")
        return

    faculty_filtered_df = faculty_df[faculty_df['faculty name'].isin(selected_faculty)]

    assigned_slots = {}
    assigned_subjects_by_year = {}
    faculty_assigned_dates = {faculty: [] for faculty in faculty_filtered_df['faculty name']}

    for year in subjects_df.columns:
        assigned_subjects_by_year[year] = []

    for exam_date in exam_dates:
        for time_slot in time_slots:
            for year in subjects_df.columns:
                subjects_list = subjects_df[year].dropna().tolist()
                subject_to_assign = None

                for subject in subjects_list:
                    if subject not in assigned_subjects_by_year[year]:
                        subject_to_assign = subject
                        break

                if subject_to_assign:
                    room_index = len(assigned_slots) % len(rooms_df)
                    room = rooms_df.iloc[room_index]
                    room_number = room.get('room number', 'N/A')

                    faculty_count = 2 if room.get('capacity', 0) > 20 else 1
                    faculty_list = []

                    # Get faculty who haven't been assigned on the current date and time
                    available_faculty = [faculty for faculty in faculty_filtered_df['faculty name'] 
                                         if time_slot not in faculty_assigned_dates[faculty]]

                    # Select faculty for the exam
                    for j in range(faculty_count):
                        if not available_faculty:
                            messagebox.showwarning("Warning", "Not enough available faculty members!")
                            return

                        faculty_index = (len(assigned_slots) + j) % len(available_faculty)
                        faculty_name = available_faculty[faculty_index]

                        if faculty_name in faculty_list:
                            continue  # Skip if faculty already in the list
                        faculty_list.append(faculty_name)

                    if faculty_list:
                        # Create a new exam entry to check for conflicts
                        new_exam = {
                            'date': exam_date,
                            'time': time_slot,
                            'room_number': room_number,
                            'faculty': faculty_list,
                            'subject': subject_to_assign
                        }

                        conflict_check_result = check_conflicts(generated_timetable, new_exam)

                        if conflict_check_result == "No conflicts":
                            assigned_slots[(exam_date, time_slot)] = subject_to_assign
                            assigned_subjects_by_year[year].append(subject_to_assign)

                            # Update faculty assigned dates
                            for faculty in faculty_list:
                                faculty_assigned_dates[faculty].append(time_slot)

                            building = room.get('building', 'N/A')
                            faculty_names = ", ".join(faculty_list)

                            # Store the new exam as a dictionary in the generated_timetable
                            row_data = {
                                'date': exam_date,
                                'time': time_slot,
                                'year': year,
                                'subject': subject_to_assign,
                                'room_number': room_number,
                                'building': building,
                                'faculty': faculty_names
                            }
                            timetable_tree.insert("", tk.END, values=(exam_date, time_slot, year, subject_to_assign, room_number, building, faculty_names))
                            generated_timetable.append(row_data)  # Append as a dictionary
                        else:
                            # If a conflict occurs, try assigning a different time slot
                            for alternate_time_slot in time_slots:
                                if alternate_time_slot == time_slot:
                                    continue  # Skip the current time slot

                                new_exam['time'] = alternate_time_slot
                                conflict_check_result = check_conflicts(generated_timetable, new_exam)

                                if conflict_check_result == "No conflicts":
                                    assigned_slots[(exam_date, alternate_time_slot)] = subject_to_assign
                                    assigned_subjects_by_year[year].append(subject_to_assign)

                                    # Update faculty assigned dates
                                    for faculty in faculty_list:
                                        faculty_assigned_dates[faculty].append(alternate_time_slot)

                                    building = room.get('building', 'N/A')
                                    faculty_names = ", ".join(faculty_list)

                                    # Store the new exam as a dictionary in the generated_timetable
                                    row_data = {
                                        'date': exam_date,
                                        'time': alternate_time_slot,
                                        'year': year,
                                        'subject': subject_to_assign,
                                        'room_number': room_number,
                                        'building': building,
                                        'faculty': faculty_names
                                    }
                                    timetable_tree.insert("", tk.END, values=(exam_date, alternate_time_slot, year, subject_to_assign, room_number, building, faculty_names))
                                    generated_timetable.append(row_data)
                                    break
                            else:
                                messagebox.showinfo("Conflict Detected", conflict_check_result)

    timetable_tree.bind("<Double-1>", lambda event: edit_cell(timetable_tree, event))

def edit_cell(tree, event):
    region = tree.identify("region", event.x, event.y)
    if region != "cell":
        return

    column = tree.identify_column(event.x)
    row_id = tree.identify_row(event.y)
    column_number = int(column.replace('#', '')) - 1

    old_value = tree.set(row_id, column_number)

    new_value = simpledialog.askstring("Edit Cell", f"Edit the value:", initialvalue=old_value)
    
    if new_value:
        tree.set(row_id, column_number, new_value)

      
        row_idx = int(row_id.strip('I')) - 1
        row_data = list(generated_timetable[row_idx])
        row_data[column_number] = new_value
        generated_timetable[row_idx] = tuple(row_data)


from docx import Document
from tkinter import messagebox, filedialog

def export_timetable():
    if not generated_timetable:
        messagebox.showwarning("Warning", "No timetable available to export.")
        return

    doc = Document()
    doc.add_heading('Exam Timetable', level=0)

    # Create the table with 1 header row and 7 columns
    table = doc.add_table(rows=1, cols=7)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Date'
    hdr_cells[1].text = 'Time'
    hdr_cells[2].text = 'Year'
    hdr_cells[3].text = 'Subject'
    hdr_cells[4].text = 'Room Number'
    hdr_cells[5].text = 'Building'
    hdr_cells[6].text = 'Faculty'

    # Populate the table with actual timetable data
    for row in generated_timetable:
        row_cells = table.add_row().cells
        
        # Check if row is a dictionary or a list
        if isinstance(row, dict):
            row_cells[0].text = str(row.get('date', ''))
            row_cells[1].text = str(row.get('time', ''))
            row_cells[2].text = str(row.get('year', ''))
            row_cells[3].text = str(row.get('subject', ''))
            row_cells[4].text = str(row.get('room_number', ''))
            row_cells[5].text = str(row.get('building', ''))
            row_cells[6].text = str(row.get('faculty', ''))
        elif isinstance(row, list) and len(row) == 7:
            for i in range(len(row)):
                row_cells[i].text = str(row[i])
        else:
            messagebox.showerror("Error", "Row format is not recognized.")

    file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Document", "*.docx")])
    if file_path:
        doc.save(file_path)
        messagebox.showinfo("Success", "Timetable exported successfully!")


root = tk.Tk()
root.title("Timetable Generator")
root.geometry("800x600")

scrollbar_y = tk.Scrollbar(root, orient=tk.VERTICAL)
scrollbar_y.pack(side=tk.RIGHT, fill=tk.Y)

scrollbar_x = tk.Scrollbar(root, orient=tk.HORIZONTAL)
scrollbar_x.pack(side=tk.BOTTOM, fill=tk.X)

subject_frame = ttk.LabelFrame(root, text="Subjects")
subject_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=10, pady=10)

room_frame = ttk.LabelFrame(root, text="Rooms")
room_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=10, pady=10)

faculty_frame = ttk.LabelFrame(root, text="Faculty")
faculty_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=10, pady=10)

content_frame = ttk.LabelFrame(root, text="Generated Timetable")
content_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=10, pady=10)

menu_bar = tk.Menu(root)
root.config(menu=menu_bar)

file_menu = tk.Menu(menu_bar, tearoff=0)
menu_bar.add_cascade(label="File", menu=file_menu)
file_menu.add_command(label="Load Subject Data", command=load_subject_data)
file_menu.add_command(label="Load Room Data", command=load_room_data)
file_menu.add_command(label="Load Faculty Data", command=load_faculty_data)
file_menu.add_separator()
file_menu.add_command(label="Exit", command=root.quit)

subject_menu = tk.Menu(menu_bar, tearoff=0)
menu_bar.add_cascade(label="Subjects", menu=subject_menu)
subject_menu.add_command(label="Add Subject", command=add_subject)
subject_menu.add_command(label="Delete Subject", command=delete_subject)

room_menu = tk.Menu(menu_bar, tearoff=0)
menu_bar.add_cascade(label="Rooms", menu=room_menu)
room_menu.add_command(label="Add Room", command=add_room)
room_menu.add_command(label="Delete Room", command=delete_room)

timetable_menu = tk.Menu(menu_bar, tearoff=0)
menu_bar.add_cascade(label="Timetable", menu=timetable_menu)
timetable_menu.add_command(label="Generate Timetable", command=generate_timetable)
timetable_menu.add_command(label="Export Timetable", command=export_timetable)

root.mainloop()
